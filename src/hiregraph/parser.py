"""File parser for resumes and job descriptions.

Supports: PDF, DOCX, Markdown, and plain text files.
"""

import os
from pathlib import Path


def parse_file(file_path: str) -> str:
    """Parse a file and return its text content.
    
    Supports:
    - .pdf (via PyMuPDF)
    - .docx (via python-docx)
    - .md, .txt, .text (read as-is)
    
    Args:
        file_path: Path to the file to parse.
        
    Returns:
        Extracted text content as a string.
        
    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file type is not supported.
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    suffix = path.suffix.lower()
    
    if suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix == ".docx":
        return _parse_docx(path)
    elif suffix in (".md", ".txt", ".text", ""):
        return _parse_text(path)
    else:
        # Try reading as plain text
        try:
            return _parse_text(path)
        except UnicodeDecodeError:
            raise ValueError(
                f"Unsupported file type: {suffix}. "
                f"Supported formats: .pdf, .docx, .md, .txt"
            )


def _parse_pdf(path: Path) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "PyMuPDF is required for PDF parsing. "
            "Install it with: pip install pymupdf"
        )
    
    doc = fitz.open(str(path))
    text_parts = []
    
    for page in doc:
        text_parts.append(page.get_text())
    
    doc.close()
    
    text = "\n".join(text_parts).strip()
    
    if not text:
        raise ValueError(f"Could not extract text from PDF: {path.name}")
    
    return text


def _parse_docx(path: Path) -> str:
    """Extract text from a Word (.docx) file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for Word document parsing. "
            "Install it with: pip install python-docx"
        )
    
    doc = Document(str(path))
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Also extract from tables (resumes sometimes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    text = "\n".join(text_parts).strip()
    
    if not text:
        raise ValueError(f"Could not extract text from DOCX: {path.name}")
    
    return text


def _parse_text(path: Path) -> str:
    """Read a plain text or markdown file."""
    return path.read_text(encoding="utf-8")


def detect_file_type(file_path: str) -> str:
    """Detect and return the file type.
    
    Returns one of: 'pdf', 'docx', 'markdown', 'text'
    """
    suffix = Path(file_path).suffix.lower()
    
    if suffix == ".pdf":
        return "pdf"
    elif suffix == ".docx":
        return "docx"
    elif suffix == ".md":
        return "markdown"
    else:
        return "text"
